from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "BirdCLEF_2026_Project_Report.docx"
EDA = ROOT / "reports" / "eda"

ACCENT = "2F7D68"
LIGHT = "E8F3EF"
GRID = "C7D6D0"
MUTED = RGBColor(90, 90, 90)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = GRID, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.color.rgb = MUTED
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, before, after in [
        ("Title", 22, 0, 12),
        ("Subtitle", 12, 0, 10),
        ("Heading 1", 16, 12, 6),
        ("Heading 2", 14, 10, 4),
        ("Heading 3", 12, 8, 3),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        if "Heading" in name or name == "Title":
            style.font.bold = True
            style.font.color.rgb = RGBColor(28, 90, 75)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "BirdCLEF+ 2026 Project Report"
    hp.style = doc.styles["Header"]
    hp.runs[0].font.name = "Arial"
    hp.runs[0].font.size = Pt(9)
    hp.runs[0].font.color.rgb = MUTED
    hp.paragraph_format.border_bottom = None

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)
    for run in fp.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    if widths is None:
        widths = [6.5 / len(headers)] * len(headers)
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = Inches(width)

    header_cells = table.rows[0].cells
    for cell, text in zip(header_cells, headers):
        cell.text = text
        set_cell_shading(cell, LIGHT)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(10)

    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 0.1


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    cell = table.cell(0, 0)
    cell.text = text
    set_cell_shading(cell, LIGHT)
    set_cell_border(cell, ACCENT, "8")
    set_cell_margins(cell, top=160, bottom=160, start=180, end=180)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(45, 75, 68)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.add_run(item)


def add_figure(doc: Document, image_name: str, caption: str, width: float = 6.2) -> None:
    path = EDA / image_name
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    inline = doc.inline_shapes[-1]._inline
    doc_pr = inline.docPr
    doc_pr.set("title", caption.split(". ", 1)[0])
    doc_pr.set("descr", caption)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.name = "Arial"
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.italic = True
    cap.runs[0].font.color.rgb = MUTED


def main() -> None:
    doc = Document()
    style_document(doc)
    add_header_footer(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("BirdCLEF+ 2026 Project Report")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Exploratory analysis, baseline modeling, and submission strategy")

    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Project", "BirdCLEF+ 2026 bioacoustic classification"],
            ["Repository", "birdclef-2026"],
            ["Primary baseline", "EfficientNet-B0 mel-spectrogram classifier"],
            ["Transfer-learning experiment", "Google Perch v2 frozen embedding probe"],
            ["Report source", "reports/PROJECT_REPORT.md and Kaggle EDA artifacts"],
        ],
        widths=[2.0, 4.5],
    )

    add_callout(
        doc,
        "Executive summary: the project establishes a reproducible Kaggle workflow, identifies the major data risks, "
        "and compares a reliable EfficientNet-B0 submission path with a stronger but operationally constrained Perch v2 probe.",
    )

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This project develops a reproducible Kaggle workflow for the BirdCLEF+ 2026 bioacoustic classification competition. "
        "The repository organizes exploratory data analysis, baseline modeling, Perch v2 transfer-learning experiments, and submission notebooks. "
        "The analysis shows 35,549 training recordings across 206 primary classes, with substantial class imbalance, uneven metadata quality, "
        "source concentration, and a distinct soundscape domain. EfficientNet-B0 reached 0.7318 validation accuracy and remains the dependable "
        "submission path. The Perch v2 probe reached 0.8403 validation accuracy, indicating strong representational value, but direct hidden-test "
        "inference is constrained by Kaggle runtime and TensorFlow/CUDA compatibility."
    )

    doc.add_heading("1. Project Objectives", level=1)
    add_numbered(
        doc,
        [
            "Establish a clean and repeatable repository structure for BirdCLEF+ 2026 experimentation.",
            "Produce exploratory analysis that identifies data quality issues, class imbalance, domain shift, and modeling implications.",
            "Implement a fast baseline model that can train and submit reliably on Kaggle.",
            "Evaluate Perch v2 embeddings as a stronger transfer-learning representation while documenting operational limitations.",
        ],
    )

    doc.add_heading("2. Repository Design", level=1)
    doc.add_paragraph(
        "The repository separates exploratory notebooks, reusable source code, model artifacts, and written reports. "
        "This design keeps Kaggle notebooks usable while avoiding notebook-only experimentation."
    )
    add_table(
        doc,
        ["Component", "Function"],
        [
            ["notebooks/01_data_eda.ipynb", "Exploratory data analysis and artifact packaging"],
            ["notebooks/02_effnet_b0_baseline.ipynb", "EfficientNet-B0 training workflow"],
            ["notebooks/03_perch_v2_probe.ipynb", "Perch v2 embedding extraction and probe training"],
            ["notebooks/04_effnet_b0_submission.ipynb", "EfficientNet-B0 submission generation"],
            ["src/birdclef2026/", "Reusable package for audio, data, model, and configuration logic"],
            ["reports/eda/", "Lightweight EDA outputs and written insight summary"],
        ],
        widths=[2.7, 3.8],
    )

    doc.add_heading("3. Dataset Summary", level=1)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Training recordings", "35,549"],
            ["Primary labels", "206"],
            ["Audio paths found", "35,549 / 35,549"],
            ["Taxonomy labels", "234 total; all 206 training labels covered"],
            ["Soundscape labels", "1,478 rows, deduplicated to 739 unique segments"],
            ["Public sample submission", "3 rows"],
        ],
        widths=[2.6, 3.9],
    )
    doc.add_paragraph(
        "The soundscape duplication result is operationally important. Clean clip metadata appears structurally safe, "
        "but soundscape prevalence and overlap analysis should be performed on deduplicated rows."
    )

    doc.add_heading("4. Exploratory Findings", level=1)
    doc.add_heading("4.1 Class Imbalance", level=2)
    add_bullets(
        doc,
        [
            "Median recordings per class: 125.",
            "Minimum and maximum recordings per class: 1 and 499.",
            "Top 10 labels account for 13.9% of recordings; top 30 account for 40.3%.",
            "Four classes are singletons.",
        ],
    )
    add_figure(doc, "class_imbalance_diagnostics.png", "Figure 1. Class-count distribution and cumulative share by ranked class.")

    doc.add_heading("4.2 Secondary Labels", level=2)
    doc.add_paragraph(
        "Secondary labels provide a noisy but useful multi-label signal. The artifacts contain 7,431 secondary-label mentions "
        "across 161 distinct labels. They are best reserved for soft-label training, mixup targets, co-occurrence priors, "
        "or confusion analysis after the primary-label baseline is stable."
    )

    doc.add_heading("4.3 Metadata Quality And Source Bias", level=2)
    add_bullets(
        doc,
        [
            "Rating 0.0 appears in 12,849 recordings.",
            "Ratings 4.0 and 5.0 together account for 14,863 recordings.",
            "Collection split: XC has 23,043 recordings and iNat has 12,506.",
            "The type field is empty for 12,975 rows.",
            "The top author contributes 2,874 recordings.",
        ],
    )
    add_figure(doc, "metadata_quality_rating.png", "Figure 2. Recording rating distribution and class-level rating variation.")

    doc.add_heading("4.4 Geography And Soundscape Domain", level=2)
    doc.add_paragraph(
        "Only 847 recordings, or 2.38%, fall inside the rough Pantanal box used in the EDA, although these rows cover 119 species. "
        "The deduplicated soundscape labels are strongly multi-label: the median segment has 4 labels, the 90th percentile has 7, "
        "and the maximum has 10. Labeled segments cluster most strongly around evening and night hours, especially 20:00-23:00."
    )
    add_figure(doc, "soundscape_overlap_time.png", "Figure 3. Soundscape label density and labeled segment distribution by hour.")

    doc.add_heading("5. Modeling Methodology", level=1)
    doc.add_heading("5.1 EfficientNet-B0 Baseline", level=2)
    doc.add_paragraph(
        "The baseline converts each recording into a 5-second mono mel-spectrogram and trains an EfficientNet-B0 classifier "
        "over 206 primary labels. The approach prioritizes fast Kaggle execution, pure PyTorch deployment, straightforward "
        "artifact packaging, reproducible fold assignment, and compatibility with competition submission constraints."
    )

    doc.add_heading("5.2 Perch v2 Probe", level=2)
    doc.add_paragraph(
        "The Perch v2 workflow uses frozen Google Perch embeddings of dimension 1,536 and trains a shallow PyTorch probe. "
        "This tests whether a bioacoustic foundation model provides stronger features than the small supervised baseline. "
        "The inspected embedding matrix has shape 35,549 x 1,536 and is not committed because it is approximately 202 MB."
    )

    doc.add_page_break()
    doc.add_heading("6. Results", level=1)
    add_table(
        doc,
        ["Model", "Representation", "Epochs", "Best validation accuracy", "Operational note"],
        [
            ["EfficientNet-B0", "5-second mel-spectrogram", "5", "0.7318", "Reliable Kaggle path"],
            ["Perch v2 probe", "Frozen 1,536-d embeddings", "10", "0.8403", "Higher validation; hidden inference constrained"],
        ],
        widths=[1.25, 1.75, 0.7, 1.2, 1.6],
    )
    doc.add_paragraph(
        "The Perch probe outperforms the EfficientNet-B0 baseline by approximately 10.9 percentage points in validation accuracy. "
        "However, EfficientNet-B0 remains the primary competition-safe artifact because it avoids TensorFlow installation issues "
        "and runs within practical hidden-test limits."
    )

    doc.add_heading("7. Submission Considerations", level=1)
    doc.add_paragraph(
        "The public sample submission has only 3 rows, so successful public dry-run execution does not guarantee hidden-test feasibility. "
        "This is especially important for heavy feature extraction workflows. EfficientNet-B0 is currently preferred for submission; "
        "Perch v2 should be used cautiously unless cached embeddings cover the evaluated rows or hidden-test runtime is proven acceptable."
    )

    doc.add_heading("8. Limitations", level=1)
    add_bullets(
        doc,
        [
            "Validation accuracy is not a complete proxy for competition score under soundscape domain shift.",
            "The EfficientNet baseline uses primary labels only and does not yet exploit secondary labels.",
            "Soundscape annotations are diagnostic rather than fully integrated into training.",
            "Perch v2 direct inference is operationally constrained on Kaggle.",
            "Metadata confounders are analyzed but not yet fully incorporated into validation or training.",
        ],
    )

    doc.add_heading("9. Recommended Next Steps", level=1)
    add_bullets(
        doc,
        [
            "Add per-class validation metrics and confusion analysis.",
            "Introduce class-aware sampling or rare-class augmentation.",
            "Add multi-crop inference for EfficientNet-B0.",
            "Explore secondary-label soft targets after the baseline remains stable.",
            "Use soundscape hour and label-overlap patterns for threshold calibration.",
            "Test knowledge distillation from Perch embeddings into a faster PyTorch model.",
            "Create validation splits that account for source, geography, and recording groups.",
        ],
    )

    doc.add_heading("10. Conclusion", level=1)
    doc.add_paragraph(
        "This project establishes a professional, reproducible BirdCLEF+ 2026 workflow with clear separation between data analysis, "
        "baseline modeling, transfer-learning experiments, and submission packaging. The EDA identifies the main risks for model "
        "generalization: class imbalance, duplicated soundscape labels, metadata quality variation, source concentration, and soundscape "
        "domain shift. EfficientNet-B0 provides a dependable submission-ready model, while the Perch v2 probe demonstrates the value "
        "of bioacoustic foundation embeddings for future improvement."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
